import os
import logging
from datetime import datetime

from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.types import FSInputFile, ReplyKeyboardRemove
from docx import Document
from docx2pdf import convert
from db_handler.db import create_connection


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

BOT_TOKEN = "—Å—é–¥–∞ –Ω–∞–¥–æ –≤—Å—Ç–∞–≤–∏—Ç—å —Ç–æ–∫–µ–Ω –±–æ—Ç–∏–∫–∞"
OUTPUT_DIR = "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()

class CertificateData:
    def __init__(self):
        self.user_id = None
        self.full_name = None
        self.birth_date = None
        self.diagnosis = None
        self.start_date = None
        self.end_date = None
        self.doctor_id = None
        self.doctor_name = None
        self.clinic_name = "–†–∞–∑—É–º–µ–¥"
        self.waiting_for = None


user_data = {}

def get_user_data(tg_id: int) -> dict:
    conn = create_connection()
    if not conn:
        return None
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT * FROM users WHERE tg_id = %s", (str(tg_id),))
            if cursor.rowcount == 0:
                return None
            columns = [desc[0] for desc in cursor.description]
            return dict(zip(columns, cursor.fetchone()))
    except Exception as e:
        logger.error(f"DB error get_user_data: {e}")
        return None
    finally:
        conn.close()

def get_last_appointment(user_id: int) -> dict:
    conn = create_connection()
    if not conn:
        return None
    try:
        with conn.cursor() as cursor:
            cursor.execute(
                "SELECT * FROM appointments WHERE user_id = %s ORDER BY appointment_date DESC LIMIT 1",
                (user_id,)
            )
            if cursor.rowcount == 0:
                return None
            columns = [desc[0] for desc in cursor.description]
            return dict(zip(columns, cursor.fetchone()))
    except Exception as e:
        logger.error(f"DB error get_last_appointment: {e}")
        return None
    finally:
        conn.close()

def get_doctor_data(doctor_id: int) -> dict:
    conn = create_connection()
    if not conn:
        return None
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT * FROM doctors WHERE doctor_id = %s", (doctor_id,))
            if cursor.rowcount == 0:
                return None
            columns = [desc[0] for desc in cursor.description]
            return dict(zip(columns, cursor.fetchone()))
    except Exception as e:
        logger.error(f"DB error get_doctor_data: {e}")
        return None
    finally:
        conn.close()

def is_valid_date(date_str: str) -> bool:
    try:
        datetime.strptime(date_str, "%d.%m.%Y")
        return True
    except ValueError:
        return False

def is_date_before(date_str1: str, date_str2: str) -> bool:
    try:
        d1 = datetime.strptime(date_str1, "%d.%m.%Y")
        d2 = datetime.strptime(date_str2, "%d.%m.%Y")
        return d1 <= d2
    except ValueError:
        return False

def is_current_year(date_str: str) -> bool:
    try:
        date = datetime.strptime(date_str, "%d.%m.%Y")
        return date.year == datetime.now().year
    except ValueError:
        return False

@dp.message(Command("start"))
async def cmd_start(message: types.Message):
    tg_id = message.from_user.id
    logger.info(f"User {tg_id} started bot")

    db_user = get_user_data(tg_id)
    if not db_user:
        await message.answer("–í—ã –Ω–µ –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω—ã –≤ —Å–∏—Å—Ç–µ–º–µ. –ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä—É–π—Ç–µ—Å—å.")
        return

    data = CertificateData()
    data.user_id = db_user['user_id']
    data.full_name = f"{db_user['first_name']} {db_user['last_name']}"
    data.birth_date = db_user.get('birth_date')

    last_appointment = get_last_appointment(data.user_id)
    if last_appointment:
        data.doctor_id = last_appointment['doctor_id']
        doctor = get_doctor_data(data.doctor_id)
        if doctor:
            data.doctor_name = f"{doctor['first_name']} {doctor['last_name']}"

    user_data[tg_id] = data

    await ask_diagnosis(message, data)

async def ask_diagnosis(message: types.Message, data: CertificateData):
    data.waiting_for = 'diagnosis'
    await message.answer(
        "–í—ã–±–µ—Ä–∏—Ç–µ –¥–∏–∞–≥–Ω–æ–∑:",
        reply_markup=types.ReplyKeyboardMarkup(
            keyboard=[[types.KeyboardButton(text=txt)] for txt in ["–û–†–í–ò", "–ì—Ä–∏–ø–ø", "–ê–Ω–≥–∏–Ω–∞", "–î—Ä—É–≥–æ–µ"]],
            resize_keyboard=True
        )
    )

@dp.message()
async def general_handler(message: types.Message):
    tg_id = message.from_user.id
    if tg_id not in user_data:
        await message.answer("–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –Ω–∞—á–Ω–∏—Ç–µ —Å–Ω–∞—á–∞–ª–∞ —Å /start.")
        return

    data = user_data[tg_id]
    text = message.text.strip()

    if data.waiting_for == 'diagnosis':
        if text not in ["–û–†–í–ò", "–ì—Ä–∏–ø–ø", "–ê–Ω–≥–∏–Ω–∞", "–î—Ä—É–≥–æ–µ"]:
            await message.answer("–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –≤—ã–±–µ—Ä–∏—Ç–µ –¥–∏–∞–≥–Ω–æ–∑ –∏–∑ —Å–ø–∏—Å–∫–∞.")
            return
        if text == "–î—Ä—É–≥–æ–µ":
            data.waiting_for = 'custom_diagnosis'
            await message.answer("–í–≤–µ–¥–∏—Ç–µ –≤–∞—à –¥–∏–∞–≥–Ω–æ–∑:", reply_markup=ReplyKeyboardRemove())
        else:
            data.diagnosis = text
            data.waiting_for = 'start_date'
            await message.answer("–í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É –Ω–∞—á–∞–ª–∞ –±–æ–ª–µ–∑–Ω–∏ (–î–î.–ú–ú.–ì–ì–ì–ì):", reply_markup=ReplyKeyboardRemove())
        return

    if data.waiting_for == 'custom_diagnosis':
        data.diagnosis = text
        data.waiting_for = 'start_date'
        await message.answer("–í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É –Ω–∞—á–∞–ª–∞ –±–æ–ª–µ–∑–Ω–∏ (–î–î.–ú–ú.–ì–ì–ì–ì):")
        return

    if data.waiting_for == 'start_date':
        if not is_valid_date(text):
            await message.answer("–ù–µ–≤–µ—Ä–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç –¥–∞—Ç—ã. –í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É –≤ —Ñ–æ—Ä–º–∞—Ç–µ –î–î.–ú–ú.–ì–ì–ì–ì.")
            return
        if not is_current_year(text):
            await message.answer(f"–î–∞—Ç–∞ –Ω–∞—á–∞–ª–∞ –¥–æ–ª–∂–Ω–∞ –±—ã—Ç—å –≤ {datetime.now().year} –≥–æ–¥—É. –ü–æ–ø—Ä–æ–±—É–π—Ç–µ —Å–Ω–æ–≤–∞.")
            return
        data.start_date = text
        data.waiting_for = 'end_date'
        await message.answer("–í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É –æ–∫–æ–Ω—á–∞–Ω–∏—è –±–æ–ª–µ–∑–Ω–∏ (–î–î.–ú–ú.–ì–ì–ì–ì):")
        return

    if data.waiting_for == 'end_date':
        if not is_valid_date(text):
            await message.answer("–ù–µ–≤–µ—Ä–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç –¥–∞—Ç—ã. –í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É –≤ —Ñ–æ—Ä–º–∞—Ç–µ –î–î.–ú–ú.–ì–ì–ì–ì.")
            return
        if not is_current_year(text):
            await message.answer(f"–î–∞—Ç–∞ –æ–∫–æ–Ω—á–∞–Ω–∏—è –¥–æ–ª–∂–Ω–∞ –±—ã—Ç—å –≤ {datetime.now().year} –≥–æ–¥—É. –ü–æ–ø—Ä–æ–±—É–π—Ç–µ —Å–Ω–æ–≤–∞.")
            return
        if not is_date_before(data.start_date, text):
            await message.answer("–î–∞—Ç–∞ –æ–∫–æ–Ω—á–∞–Ω–∏—è –Ω–µ –º–æ–∂–µ—Ç –±—ã—Ç—å —Ä–∞–Ω—å—à–µ –¥–∞—Ç—ã –Ω–∞—á–∞–ª–∞. –í–≤–µ–¥–∏—Ç–µ –¥–∞—Ç—É –æ–∫–æ–Ω—á–∞–Ω–∏—è –∑–∞–Ω–æ–≤–æ:")
            return
        data.end_date = text
        await generate_certificate(message, data)
        return

    await message.answer("–ù–µ–∏–∑–≤–µ—Å—Ç–Ω–∞—è –∫–æ–º–∞–Ω–¥–∞. –ü–æ–ø—Ä–æ–±—É–π—Ç–µ —Å–Ω–æ–≤–∞ –∏–ª–∏ –Ω–∞—á–Ω–∏—Ç–µ —Å /start.")

async def generate_certificate(message: types.Message, data: CertificateData):
    tg_id = message.from_user.id
    docx_path = os.path.join(OUTPUT_DIR, f"certificate_{tg_id}.docx")
    pdf_path = os.path.join(OUTPUT_DIR, f"certificate_{tg_id}.pdf")

    doc = Document()
    doc.add_heading("–ú–µ–¥–∏—Ü–∏–Ω—Å–∫–∞—è —Å–ø—Ä–∞–≤–∫–∞", level=1)
    doc.add_paragraph(f"–í—ã–¥–∞–Ω–∞: {data.full_name}")
    doc.add_paragraph(f"–î–∞—Ç–∞ —Ä–æ–∂–¥–µ–Ω–∏—è: {data.birth_date}")
    doc.add_paragraph(f"–î–∏–∞–≥–Ω–æ–∑: {data.diagnosis}")
    doc.add_paragraph(f"–ü–µ—Ä–∏–æ–¥ –±–æ–ª–µ–∑–Ω–∏: —Å {data.start_date} –ø–æ {data.end_date}")
    doc.add_paragraph("–°–ø—Ä–∞–≤–∫–∞ –≤—ã–¥–∞–Ω–∞ –¥–ª—è –ø—Ä–µ–¥–æ—Å—Ç–∞–≤–ª–µ–Ω–∏—è –ø–æ –º–µ—Å—Ç—É —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è.")
    doc.add_paragraph(f"–í—Ä–∞—á: {data.doctor_name or '–ù–µ —É–∫–∞–∑–∞–Ω'}")
    doc.add_paragraph(f"–ú–µ–¥–∏—Ü–∏–Ω—Å–∫–æ–µ —É—á—Ä–µ–∂–¥–µ–Ω–∏–µ: {data.clinic_name}")
    doc.add_paragraph(f"–î–∞—Ç–∞ –≤—ã–¥–∞—á–∏: {datetime.now().strftime('%d.%m.%Y')}")

    doc.save(docx_path)
    convert(docx_path, pdf_path)

    await message.answer("–í–∞—à–∞ —Å–ø—Ä–∞–≤–∫–∞ –≥–æ—Ç–æ–≤–∞!")
    await message.answer_document(FSInputFile(pdf_path))

    os.remove(docx_path)
    os.remove(pdf_path)
    del user_data[tg_id]


@dp.message(F.text.startswith("/symptom"))
async def handle_symptom_with_advice(message: Message):
    user_input = message.text.replace("/symptom", "").strip()

    specialization = model_predict(user_input)  # todo: –¥–æ–±–∞–≤–∏—Ç—å –º–æ–¥–µ–ª—å

    if not specialization:
        await message.answer("üòï –ù–µ —É–¥–∞–ª–æ—Å—å –æ–ø—Ä–µ–¥–µ–ª–∏—Ç—å —Å–ø–µ—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—é –ø–æ –≤–≤–µ–¥—ë–Ω–Ω—ã–º —Å–∏–º–ø—Ç–æ–º–∞–º.")
        return

    try:
        conn = create_connection()
        if not conn:
            return None
            
        with conn.cursor() as cursor:
            tips = cursor.execute(
                "SELECT title, content FROM MedicalTips WHERE category = $1",
                specialization
            )
        conn.close()
    except Exception as e:
        await message.answer(f"‚ö†Ô∏è –û—à–∏–±–∫–∞ –ø—Ä–∏ –ø–æ–¥–∫–ª—é—á–µ–Ω–∏–∏ –∫ –±–∞–∑–µ –¥–∞–Ω–Ω—ã—Ö: {e}")
        return

    if not tips:
        await message.answer(f"–°–æ–≤–µ—Ç—ã –ø–æ —Å–ø–µ—Ü–∏–∞–ª–∏–∑–∞—Ü–∏–∏ ¬´{specialization}¬ª –ø–æ–∫–∞ –æ—Ç—Å—É—Ç—Å—Ç–≤—É—é—Ç.")
        return

    tip = random.choice(tips)
    response = f"üí° *{tip['title']}*\n\n{tip['content']}"
    await message.answer(response, parse_mode="Markdown")

if __name__ == "__main__":
    dp.run_polling(bot)
