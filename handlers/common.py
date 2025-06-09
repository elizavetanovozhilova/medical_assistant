from aiogram import Router, types
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton, ReplyKeyboardRemove, FSInputFile
from keyboards.inline import get_main_menu, get_gender_keyboard, get_help_back_keyboard, get_medcard_keyboard,get_dates_keyboard,get_times_keyboard,get_doctors_keyboard
from db_handler.db import (
    check_auth,
    register_user,
    get_doctors,
    generate_available_times,
    generate_available_dates,
    create_appointment,
    get_doctor_info,
    get_doctors_by_specialization,
    get_user_diagnoses,
    get_doctor_data,
    get_last_appointment,
    get_user_data,
    get_last_diagnosis,
    get_user_appointments
)
from keyboards.reply import get_menu_reply_keyboard
from utils import get_user_data, get_last_appointment, get_doctor_data
from model import get_doctor, predict_intent, get_date

from docx import Document
from docx2pdf import convert
from datetime import datetime
import os
import logging

router = Router()

OUTPUT_DIR = "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)
logger = logging.getLogger(__name__)
user_data = {}


class CertificateData:
    def __init__(self):
        self.user_id = None
        self.full_name = None
        self.birth_date = None
        self.diagnosis = None
        self.start_date = None
        self.end_date = None
        self.doctor_id = None
        self.doctor_name = None
        self.clinic_name = "Разумед"
        self.waiting_for = None


class AppointmentStates(StatesGroup):
    waiting_for_doctor = State()
    waiting_for_date = State()
    waiting_for_time = State()


class RegistrationStates(StatesGroup):
    waiting_for_first_name = State()
    waiting_for_last_name = State()
    waiting_for_gender = State()
    waiting_for_birth_date = State()
    waiting_for_phone = State()


class RecommendationStates(StatesGroup):
    waiting_for_symptoms = State()


class MenuState(StatesGroup):
    waiting_for_input = State()


class CertificateStates(StatesGroup):
    waiting_for_start_date = State()
    waiting_for_end_date = State()


def is_valid_date(date_str: str) -> bool:
    try:
        datetime.strptime(date_str, "%d.%m.%Y")
        return True
    except ValueError:
        return False


def is_date_before(date_str1: str, date_str2: str) -> bool:
    try:
        d1 = datetime.strptime(date_str1, "%d.%m.%Y")
        d2 = datetime.strptime(date_str2, "%d.%m.%Y")
        return d1 <= d2
    except ValueError:
        return False


def is_current_year(date_str: str) -> bool:
    try:
        date = datetime.strptime(date_str, "%d.%m.%Y")
        return date.year == datetime.now().year
    except ValueError:
        return False


@router.message(Command("start"))
async def cmd_start(message: types.Message, state: FSMContext):
    user_id = str(message.from_user.id)

    if check_auth(user_id):
        await message.answer(
            "С возвращением! Главное меню:",
            reply_markup=get_menu_reply_keyboard()
        )
        await message.answer(
            "Выберите действие или напишите, что хотите сделать:",
            reply_markup=get_main_menu()
        )
    else:
        await message.answer("Введите ваше имя:", reply_markup=get_menu_reply_keyboard())
        await state.set_state(RegistrationStates.waiting_for_first_name)


@router.message(RegistrationStates.waiting_for_first_name)
async def process_first_name(message: types.Message, state: FSMContext):
    await state.update_data(first_name=message.text)
    await message.answer("Введите вашу фамилию:")
    await state.set_state(RegistrationStates.waiting_for_last_name)


@router.message(RegistrationStates.waiting_for_last_name)
async def process_last_name(message: types.Message, state: FSMContext):
    await state.update_data(last_name=message.text)
    await message.answer("Укажите ваш пол:", reply_markup=get_gender_keyboard())
    await state.set_state(RegistrationStates.waiting_for_gender)


@router.callback_query(RegistrationStates.waiting_for_gender)
async def process_gender(callback: types.CallbackQuery, state: FSMContext):
    gender = 'М' if callback.data == 'gender_male' else 'Ж'
    await state.update_data(gender=gender)
    await callback.message.edit_text(f"Пол: {gender}")
    await callback.message.answer("Введите вашу дату рождения (ДД.ММ.ГГГГ):")
    await state.set_state(RegistrationStates.waiting_for_birth_date)


@router.message(RegistrationStates.waiting_for_birth_date)
async def process_birth_date(message: types.Message, state: FSMContext):
    birth_date = message.text.strip()

    if not is_valid_date(birth_date):
        await message.answer("Неверный формат даты. Введите дату в формате ДД.ММ.ГГГГ.")
        return

    await state.update_data(birth_date=birth_date)
    await message.answer("Введите ваш номер телефона:")
    await state.set_state(RegistrationStates.waiting_for_phone)


@router.message(RegistrationStates.waiting_for_phone)
async def process_phone(message: types.Message, state: FSMContext):
    user_data = await state.get_data()
    user = message.from_user

    success = register_user(
        tg_id=str(user.id),
        username=user.username,
        first_name=user_data['first_name'],
        last_name=user_data['last_name'],
        gender=user_data['gender'],
        phone=message.text,
        birth_date=user_data['birth_date']
    )

    if success:
        await state.clear()
        await message.answer(
            "Регистрация успешно завершена!",
            reply_markup=get_menu_reply_keyboard()
        )
        await state.set_state(MenuState.waiting_for_input)
        await message.answer("🏠Главное меню:", reply_markup=get_main_menu())
    else:
        await message.answer("Ошибка регистрации. Попробуйте позже.")


@router.callback_query(lambda c: c.data == "main_menu")
async def process_main_menu(callback: types.CallbackQuery, state: FSMContext):
    await state.set_state(MenuState.waiting_for_input)
    await callback.message.answer(
        "🏠Главное меню:",
        reply_markup=get_main_menu()
    )


@router.message(lambda msg: msg.text == "Меню")
async def handle_menu_button(message: types.Message, state: FSMContext):
    await state.set_state(MenuState.waiting_for_input)
    await message.answer(
        "🏠Главное меню:",
        reply_markup=get_main_menu()
    )


@router.message(MenuState.waiting_for_input)
async def main_menu_text_handler(message: types.Message, state: FSMContext):
    intent = predict_intent(message.text)

    if intent == "рекомендация":
        await ask_for_symptoms(message, state)

    elif intent == "запись":
        await handler_appointment(message, state)

    elif intent == 'оставить_отзыв':
        pass
        # await функция для оставления отзыва

    elif intent == 'читать_отзыв':
        pass
        # await функция для чтения отзывов

    elif intent == "справка":
        tg_id = str(message.from_user.id)
        data = user_data.get(tg_id)
        if not data:
            await start_certificate_process(message, state)
            return
        await ask_for_symptoms(message, state)

    elif intent == 'профиль':
        await handle_medcard_profile(message, state)

    elif intent == 'приемы':
        await handle_medcard_appointments(message, state)

    elif intent == 'диагнозы':
        await handle_medcard_diagnoses(message, state)

    elif intent == 'график':
        await message.answer(
            "🕒 *График работы:*\n\n"
            "📅 *Понедельник – Пятница:* 07:30 – 20:00\n"
            "📅 *Суббота:* 08:00 – 17:00\n"
            "📅 *Воскресенье:* 08:00 – 15:00\n\n"
            "🔹 *График в праздничные дни* — уточняйте заранее.\n"
            "⛔ *Ночного приёма нет.*",
            parse_mode="Markdown"
        )

    elif intent == 'адреса':
        await message.answer(
            "🏥 *Наши филиалы:*\n\n"
            "📍 ул. Ю. Фучика, 53а\n"
            "📍 ул. Ак. Глушко, д. 15а\n"
            "📍 ул. Беломорская, д. 6",
            parse_mode="Markdown"
        )

    elif intent == 'специальности':
        await message.answer(
            "🏥 Наши специальности:\n\n"
            "• Гастроэнтерология\n"
            "• Гематология\n"
            "• Гинекология\n"
            "• Дерматовенерология\n"
            "• Кардиология\n"
            "• Маммология\n"
            "• Неврология\n"
            "• Нутрициология\n"
            "• Отоларингология\n"
            "• Офтальмология\n"
            "• Проктология\n"
            "• Психотерапия\n"
            "• Сосудистая хирургия\n"
            "• Урология\n"
            "• Хирургия\n"
            "• Эндокринология",
            parse_mode="Markdown"
        )

    elif intent == 'поддержка':
        await message.answer(
            "Если у вас возникли технические вопросы или нужна помощь, с удовольствием вам ответит наш специалист "
            "— @fliwoll. Пожалуйста, обращайтесь!",
            parse_mode="Markdown"
        )

    elif intent == 'функционал':
        await handler_help(message, state)

    else:
        await message.answer(
            "❓ Извините, я пока не понял ваш запрос.\nЕсли вы не нашли нужную "
            "функцию в меню или у вас возник вопрос — пожалуйста, напишите в поддержку: "
            "@fliwoll. Мы обязательно поможем!"
        )


# БЛОК ЗАПИСЬ НА ПРИЕМ
@router.callback_query(lambda c: c.data == "appointment")
async def start_appointment(callback: types.CallbackQuery, state: FSMContext):
    await handler_appointment(callback.message, state)


async def handler_appointment(message: types.Message, state: FSMContext):
    await state.clear()
    doctors = get_doctors()
    if not doctors:
        await message.answer("В данный момент нет доступных врачей для записи.")
        return

    await message.answer(
        "Выберите врача:",
        reply_markup=get_doctors_keyboard(doctors)
    )
    await state.set_state(AppointmentStates.waiting_for_doctor)


@router.callback_query(AppointmentStates.waiting_for_doctor, lambda c: c.data.startswith("doctor_"))
async def process_doctor_selection(callback: types.CallbackQuery, state: FSMContext):
    doctor_id = int(callback.data.split("_")[1])
    await state.update_data(doctor_id=doctor_id)

    available_dates = generate_available_dates(doctor_id)
    if not available_dates:
        await callback.message.answer("Нет доступных дат для записи к этому врачу.")
        await state.clear()
        return

    await callback.message.answer(
        "Выберите дату приема:",
        reply_markup=get_dates_keyboard(available_dates)
    )
    await state.set_state(AppointmentStates.waiting_for_date)
    await callback.answer()

@router.callback_query(AppointmentStates.waiting_for_date, lambda c: c.data.startswith("date_"))
async def process_date_selection(callback: types.CallbackQuery, state: FSMContext):
    date_str = callback.data.split("_")[1]
    data = await state.get_data()
    doctor_id = data['doctor_id']

    try:
        selected_date = datetime.strptime(date_str, "%Y-%m-%d").date()
    except ValueError:
        await callback.message.answer("Неверный формат даты.")
        await state.clear()
        return

    available_times = generate_available_times(doctor_id, selected_date)
    if not available_times:
        await callback.message.answer("Нет доступного времени для записи на выбранную дату.")
        await state.clear()
        return

    await state.update_data(date=date_str)
    await callback.message.answer(
        "Выберите время приема:",
        reply_markup=get_times_keyboard([time.strftime("%H:%M") for time in available_times])
    )
    await state.set_state(AppointmentStates.waiting_for_time)
    await callback.answer()


@router.callback_query(AppointmentStates.waiting_for_time, lambda c: c.data.startswith("time_"))
async def process_time_selection(callback: types.CallbackQuery, state: FSMContext):
    time_str = callback.data.split("_")[1]
    data = await state.get_data()

    try:
        doctor_id = data['doctor_id']
        date_str = data['date']
        tg_id = callback.from_user.id

        if create_appointment(tg_id, doctor_id, date_str, time_str):
            doctor_info = get_doctor_info(doctor_id)
            if doctor_info:
                doctor_name = f"{doctor_info['first_name']} {doctor_info['last_name']} ({doctor_info['specialization']})"
            else:
                doctor_name = "врачу"

            confirmation_message = (
                f"✅ Вы успешно записаны на прием!\n\n"
                f"👨‍⚕️ Врач: {doctor_name}\n"
                f"📅 Дата: {date_str.replace('-', '.')}\n"
                f"⏰ Время: {time_str}\n\n"
                f"Ждем вас в указанное время!"
            )

            await callback.message.answer(confirmation_message)
        else:
            await callback.message.answer("❌ Произошла ошибка при создании записи. Попробуйте позже.")

    except Exception as e:
        print(f"Error processing time selection: {e}")
        await callback.message.answer("❌ Произошла ошибка. Попробуйте еще раз.")
    finally:
        await state.clear()
        await callback.answer()


# БЛОК РЕКОМЕНДАЦИЙ
@router.callback_query(lambda c: c.data == "recommendation")
async def start_recommendation(callback: types.CallbackQuery, state: FSMContext):
    await ask_for_symptoms(callback.message, state)


async def ask_for_symptoms(message: types.Message, state: FSMContext):
    await state.clear()
    await message.answer("Пожалуйста, опишите ваши симптомы:")
    await state.set_state(RecommendationStates.waiting_for_symptoms)


@router.message(RecommendationStates.waiting_for_symptoms)
async def process_symptoms(message: types.Message, state: FSMContext):
    symptoms = message.text
    if not symptoms or not isinstance(symptoms, str):
        await message.answer("Ошибка: описание симптомов отсутствует или неверного формата.")
        return

    doctor_specialization = get_doctor(symptoms)
    await state.update_data(predicted_doctor=doctor_specialization)

    doctors = get_doctors_by_specialization(doctor_specialization)

    if doctors:
        text = (
            f"Согласно вашим симптомам: {symptoms}\n\n"
            f"Рекомендуем вам обратиться к врачу по специальности: *{doctor_specialization}*.\n\n"
            "Вы можете записаться к одному из доступных специалистов:"
        )

        keyboard = InlineKeyboardMarkup(inline_keyboard=[])

        for doctor in doctors:
            doctor_id, first_name, last_name, spec = doctor
            keyboard.inline_keyboard.append([
                InlineKeyboardButton(
                    text=f"{first_name} {last_name} ({spec})",
                    callback_data=f"rec_doctor_{doctor_id}"
                )
            ])

        keyboard.inline_keyboard.append([
            InlineKeyboardButton(text="🔙 Назад", callback_data="main_menu")
        ])

        await message.answer(
            text,
            reply_markup=keyboard,
            parse_mode="Markdown"
        )
    else:
        await message.answer(
            f"Согласно вашим симптомам: {symptoms}\n\n"
            f"Рекомендуем вам обратиться к врачу по специальности: *{doctor_specialization}*.\n\n"
            "К сожалению, сейчас нет доступных специалистов для записи.",
            parse_mode="Markdown"
        )

    await state.clear()


@router.callback_query(lambda c: c.data.startswith("rec_doctor_"))
async def process_recommended_doctor(callback: types.CallbackQuery, state: FSMContext):
    doctor_id = int(callback.data.split("_")[2])
    await state.update_data(doctor_id=doctor_id)

    available_dates = generate_available_dates(doctor_id)
    if not available_dates:
        await callback.message.answer("Нет доступных дат для записи к этому врачу.")
        await state.clear()
        return

    await callback.message.answer(
        "Выберите дату приема:",
        reply_markup=get_dates_keyboard(available_dates)
    )
    await state.set_state(AppointmentStates.waiting_for_date)
    await callback.answer()


# БЛОК МЕДКАРТЫ
@router.callback_query(lambda c: c.data == "medical_card")
async def process_help_callback(callback: types.CallbackQuery, state: FSMContext):
    await state.clear()
    await callback.message.answer(
        "Что вы хотите узнать?",
        reply_markup=get_medcard_keyboard()
    )


@router.callback_query(lambda c: c.data == "medcard_profile")
async def start_medcard_profile(callback: types.CallbackQuery, state: FSMContext):
    await handle_medcard_profile(callback, state)


async def handle_medcard_profile(callback: types.CallbackQuery, state: FSMContext):
    await state.clear()
    tg_id = str(callback.from_user.id)

    user = get_user_data(tg_id)
    if not user:
        await callback.message.answer("Профиль не найден.")
        return

    text = (
        f"👤 Профиль:\n\n"
        f"Имя: {user['first_name']}\n"
        f"Фамилия: {user['last_name']}\n"
        f"Пол: {user['gender']}\n"
        f"Телефон: {user['phone']}\n"
        f"Дата рождения: {user.get('birth_date', 'не указана')}"
    )
    await callback.message.answer(text)


@router.callback_query(lambda c: c.data == "medcard_appointments")
async def start_medcard_appointments(callback: types.CallbackQuery, state: FSMContext):
    await handle_medcard_appointments(callback, state)


async def handle_medcard_appointments(callback: types.CallbackQuery, state: FSMContext):
    await state.clear()
    tg_id = str(callback.from_user.id)
    user = get_user_data(tg_id)
    if not user:
        await callback.answer("Профиль не найден.")
        return

    appointments = get_user_appointments(user['user_id'])
    if not appointments:
        await callback.answer()
        await callback.message.answer("У вас нет запланированных приемов.")
        return

    text = "📅 Ваши последние приемы:\n\n"
    for appt in appointments:
        text += (
            f"Дата: {appt['date']}\n"
            f"Время: {appt['time']}\n"
            f"Врач: {appt['doctor_name']}\n"
            f"Специальность: {appt['specialization']}\n\n"
        )

    await callback.answer()
    await callback.message.answer(text)


@router.callback_query(lambda c: c.data == "medcard_diagnoses")
async def start_medcard_diagnoses(callback: types.CallbackQuery, state: FSMContext):
    await handle_medcard_diagnoses(callback, state)


async def handle_medcard_diagnoses(callback: types.CallbackQuery, state: FSMContext):
    await state.clear()
    tg_id = str(callback.from_user.id)
    user = get_user_data(tg_id)
    if not user:
        await callback.answer("Профиль не найден.")
        return

    diagnoses = get_user_diagnoses(user['user_id'])
    if not diagnoses:
        await callback.answer()
        await callback.message.answer("Диагнозов не найдено.")
        return

    text = "💊 Ваши диагнозы:\n\n"
    for diagnosis in diagnoses:
        text += f"{diagnosis['name']} - {diagnosis['date']}\n"

    await callback.answer()
    await callback.message.answer(text)



# БЛОК СПРАВКИ


@router.callback_query(lambda c: c.data == "reference")
async def start_certificate(callback: types.CallbackQuery, state: FSMContext):
    await process_certificate_start(callback, state)
    await callback.answer()


async def process_certificate_start(callback: types.CallbackQuery, state: FSMContext):
    tg_id = str(callback.from_user.id)
    user = get_user_data(tg_id)

    if not user:
        await callback.message.answer("Вы не зарегистрированы в системе.")
        return

    #last_diagnosis = get_last_diagnosis(user['user_id'])
    last_diagnosis = {'name':'ОРВИ', 'date': '07.06.2025'}
    if not last_diagnosis:
        await callback.message.answer("У вас нет сохраненных диагнозов.")
        return

    await state.update_data(
        user_id=user['user_id'],
        full_name=f"{user['first_name']} {user['last_name']}",
        birth_date=user.get('birth_date'),
        diagnosis=last_diagnosis['name'],
        diagnosis_date=last_diagnosis['date']
    )

    last_appointment = get_last_appointment(user['user_id'])
    if last_appointment:
        doctor = get_doctor_data(last_appointment['doctor_id'])
        if doctor:
            await state.update_data(
                doctor_name=f"{doctor['first_name']} {doctor['last_name']}"
            )

    await callback.answer(
        f"Ваш последний диагноз: {last_diagnosis['name']} (от {last_diagnosis['date']})\n"
        "Введите дату начала болезни (ДД.ММ.ГГГГ):"
    )
    await state.set_state(CertificateStates.waiting_for_start_date)


@router.message(CertificateStates.waiting_for_start_date)
async def process_start_date(message: types.Message, state: FSMContext):
    date_str = message.text.strip()

    if not is_valid_date(date_str):
        await message.answer("Неверный формат даты. Введите дату в формате ДД.ММ.ГГГГ.")
        return

    if not is_current_year(date_str):
        await message.answer(f"Дата начала должна быть в {datetime.now().year} году.")
        return

    await state.update_data(start_date=date_str)
    await message.answer("Введите дату окончания болезни (ДД.ММ.ГГГГ):")
    await state.set_state(CertificateStates.waiting_for_end_date)


@router.message(CertificateStates.waiting_for_end_date)
async def process_end_date(message: types.Message, state: FSMContext):
    data = await state.get_data()
    end_date_str = message.text.strip()

    if not is_valid_date(end_date_str):
        await message.answer("Неверный формат даты. Введите дату в формате ДД.ММ.ГГГГ.")
        return

    if not is_current_year(end_date_str):
        await message.answer(f"Дата окончания должна быть в {datetime.now().year} году.")
        return

    if not is_date_before(data['start_date'], end_date_str):
        await message.answer("Дата окончания не может быть раньше даты начала.")
        return

    await state.update_data(end_date=end_date_str)
    await generate_certificate(message, await state.get_data())
    await state.clear()


async def generate_certificate(message: types.Message, data: dict):
    tg_id = str(message.from_user.id)
    docx_path = os.path.join(OUTPUT_DIR, f"certificate_{tg_id}.docx")
    pdf_path = os.path.join(OUTPUT_DIR, f"certificate_{tg_id}.pdf")

    doc = Document()
    doc.add_heading("Медицинская справка", level=1)
    doc.add_paragraph(f"Выдана: {data['full_name']}")
    doc.add_paragraph(f"Дата рождения: {data['birth_date']}")
    doc.add_paragraph(f"Диагноз: {data['diagnosis']}")
    doc.add_paragraph(f"Период болезни: с {data['start_date']} по {data['end_date']}")
    doc.add_paragraph("Справка выдана для предоставления по месту требования.")
    doc.add_paragraph(f"Врач: {data.get('doctor_name', 'Не указан')}")
    doc.add_paragraph(f"Медицинское учреждение: Разумед")
    doc.add_paragraph(f"Дата выдачи: {datetime.now().strftime('%d.%m.%Y')}")

    doc.save(docx_path)
    convert(docx_path, pdf_path)

    await message.answer("Ваша справка готова!")
    await message.answer_document(FSInputFile(pdf_path))
    await message.answer("🏠 Главное меню:", reply_markup=get_main_menu())

    os.remove(docx_path)
    os.remove(pdf_path)


@router.message()
async def handler_certificate(message: types.Message, state: FSMContext):
    tg_id = str(message.from_user.id)
    data = user_data.get(tg_id)
    if not data:
        await process_certificate_start(message, state)
        return

    text = message.text.strip()

    if data.waiting_for == 'start_date':
        text = get_date(text)
        if not is_valid_date(text):
            await message.answer("Неверный формат даты. Введите дату в формате ДД.ММ.ГГГГ.")
            return
        if not is_current_year(text):
            await message.answer(f"Дата начала должна быть в {datetime.now().year} году.")
            return
        data.start_date = text
        data.waiting_for = 'end_date'
        await message.answer("Введите дату окончания болезни (ДД.ММ.ГГГГ):", reply_markup=ReplyKeyboardRemove())
        return

    if data.waiting_for == 'end_date':
        text = get_date(text)
        if not is_valid_date(text):
            await message.answer("Неверный формат даты. Введите дату в формате ДД.ММ.ГГГГ.")
            return
        if not is_current_year(text):
            await message.answer(f"Дата окончания должна быть в {datetime.now().year} году.")
            return
        if not is_date_before(data.start_date, text):
            await message.answer("Дата окончания не может быть раньше даты начала.")
            return
        data.end_date = text
        await generate_certificate(message, data, state)
        user_data.pop(tg_id, None)
        return

    if text == "🔙 Назад":
        await state.set_state(MenuState.waiting_for_input)
        await message.answer("🏠 Главное меню:", reply_markup=get_main_menu())
        user_data.pop(tg_id, None)
        return


# БЛОК ПОМОЩИ
@router.callback_query(lambda c: c.data == "help")
async def start_help(callback: types.CallbackQuery, state: FSMContext):
    await handler_help(callback.message, state)


async def handler_help(message: types.Message, state: FSMContext):
    await state.clear()
    await message.answer(
        "❓ Помощь\n\n"
        "📌 Основные функции:\n\n"
        "• 📝 Записаться к врачу: выбор специалиста, даты и времени.\n"
        "• 💊 Рекомендация: подбор врача по симптомам.\n"
        "• 🩺 Мед карта: ваши данные\n"
        "• 📄 Справка: справки, налоги, история посещений.\n\n"
        "⚙️ Дополнительно:\n\n"
        "• Напоминания ⏰\n"
        "• Изменение данных 👤\n"
        "• FAQ ❓\n\n"
        "🔄 Навигация:\n\n"
        "• «Назад» — вернуться в предыдущее меню.\n"
        "• «Главное меню» — к основным разделам.\n"
        "📩 Поддержка: @fliwoll\n\n"
        "Кнопки:\n\n"
        "• 🔙 Назад\n"
        "• 🏠 Главное меню",
        reply_markup=get_help_back_keyboard()
    )